#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Document processor for handling Word documents and Excel sheets.
"""

import os
import logging
import asyncio
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import tempfile
import base64

# Document processing libraries
from docx import Document
import pandas as pd
import openpyxl
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processor for Word documents and Excel sheets."""
    
    def __init__(self, config: Dict[str, Any]):
        """Initialize the document processor.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config
        self.temp_dir = tempfile.gettempdir()
        
        logger.info("Document processor initialized")
    
    async def process(self, task_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process a document (Word or Excel).
        
        Args:
            task_data: Task data containing file information
            
        Returns:
            Dictionary containing processing results
        """
        try:
            file_path = task_data.get("file_path")
            file_content = task_data.get("file_content")  # Base64 encoded
            file_name = task_data.get("file_name", "")
            file_type = task_data.get("file_type", "").lower()
            
            # Determine file type from extension if not provided
            if not file_type and file_name:
                file_extension = os.path.splitext(file_name)[1].lower()
                if file_extension in ['.docx', '.doc']:
                    file_type = 'word'
                elif file_extension in ['.xlsx', '.xls']:
                    file_type = 'excel'
            
            # Handle file content (either path or base64 content)
            temp_file_path = None
            if file_content:
                # Decode base64 content and save to temp file
                try:
                    file_data = base64.b64decode(file_content)
                    temp_file_path = os.path.join(self.temp_dir, f"temp_{datetime.now().timestamp()}_{file_name}")
                    with open(temp_file_path, 'wb') as f:
                        f.write(file_data)
                    file_path = temp_file_path
                except Exception as e:
                    logger.error(f"Error decoding file content: {e}")
                    return {"error": f"Failed to decode file content: {str(e)}"}
            
            if not file_path or not os.path.exists(file_path):
                return {"error": "File path not provided or file does not exist"}
            
            # Process based on file type
            result = {}
            if file_type == 'word':
                result = await self._process_word_document(file_path, file_name)
            elif file_type == 'excel':
                result = await self._process_excel_document(file_path, file_name)
            else:
                result = {"error": f"Unsupported file type: {file_type}"}
            
            # Clean up temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to remove temp file {temp_file_path}: {e}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            return {"error": f"Document processing failed: {str(e)}"}
    
    async def _process_word_document(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process a Word document.
        
        Args:
            file_path: Path to the Word document
            file_name: Original file name
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract text content
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append({
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract headers and footers
            headers = []
            footers = []
            for section in doc.sections:
                if section.header:
                    header_text = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text.append(paragraph.text.strip())
                    if header_text:
                        headers.append('\n'.join(header_text))
                
                if section.footer:
                    footer_text = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text.append(paragraph.text.strip())
                    if footer_text:
                        footers.append('\n'.join(footer_text))
            
            # Extract document properties
            properties = {}
            if doc.core_properties:
                core_props = doc.core_properties
                properties.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                })
            
            # Combine all text for full content
            full_text = '\n'.join([p['text'] for p in paragraphs])
            
            # Count statistics
            word_count = len(full_text.split())
            char_count = len(full_text)
            paragraph_count = len(paragraphs)
            table_count = len(tables)
            
            return {
                'file_name': file_name,
                'file_type': 'word',
                'content': {
                    'full_text': full_text,
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'headers': headers,
                    'footers': footers
                },
                'properties': properties,
                'statistics': {
                    'word_count': word_count,
                    'character_count': char_count,
                    'paragraph_count': paragraph_count,
                    'table_count': table_count
                },
                'processed_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            return {"error": f"Failed to process Word document: {str(e)}"}
    
    async def _process_excel_document(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Process an Excel document.
        
        Args:
            file_path: Path to the Excel document
            file_name: Original file name
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            # Load workbook
            workbook = load_workbook(file_path, data_only=True)
            
            # Process each worksheet
            worksheets = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract data using pandas for better handling
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
                
                # Convert to list of lists, handling NaN values
                data = df.fillna('').values.tolist()
                
                # Extract sheet metadata
                sheet_info = {
                    'name': sheet_name,
                    'data': data,
                    'dimensions': {
                        'rows': len(data),
                        'columns': len(data[0]) if data else 0
                    }
                }
                
                # Try to identify headers (first row with non-empty values)
                headers = []
                if data:
                    first_row = data[0]
                    if any(str(cell).strip() for cell in first_row):
                        headers = [str(cell).strip() for cell in first_row]
                        sheet_info['headers'] = headers
                
                # Extract summary statistics for numeric columns
                try:
                    numeric_df = pd.read_excel(file_path, sheet_name=sheet_name)
                    numeric_columns = numeric_df.select_dtypes(include=['number']).columns.tolist()
                    
                    if numeric_columns:
                        statistics = {}
                        for col in numeric_columns:
                            col_data = numeric_df[col].dropna()
                            if not col_data.empty:
                                statistics[col] = {
                                    'count': len(col_data),
                                    'mean': float(col_data.mean()),
                                    'std': float(col_data.std()),
                                    'min': float(col_data.min()),
                                    'max': float(col_data.max()),
                                    'sum': float(col_data.sum())
                                }
                        
                        if statistics:
                            sheet_info['statistics'] = statistics
                            
                except Exception as e:
                    logger.warning(f"Could not extract statistics for sheet {sheet_name}: {e}")
                
                worksheets.append(sheet_info)
            
            # Extract workbook properties
            properties = {}
            if hasattr(workbook, 'properties') and workbook.properties:
                props = workbook.properties
                properties.update({
                    'title': getattr(props, 'title', '') or '',
                    'author': getattr(props, 'creator', '') or '',
                    'subject': getattr(props, 'subject', '') or '',
                    'keywords': getattr(props, 'keywords', '') or '',
                    'comments': getattr(props, 'description', '') or '',
                    'created': getattr(props, 'created', ''),
                    'modified': getattr(props, 'modified', ''),
                    'last_modified_by': getattr(props, 'lastModifiedBy', '') or ''
                })
            
            # Convert datetime objects to strings
            for key, value in properties.items():
                if hasattr(value, 'isoformat'):
                    properties[key] = value.isoformat()
            
            # Calculate overall statistics
            total_rows = sum(ws['dimensions']['rows'] for ws in worksheets)
            total_columns = sum(ws['dimensions']['columns'] for ws in worksheets)
            sheet_count = len(worksheets)
            
            return {
                'file_name': file_name,
                'file_type': 'excel',
                'content': {
                    'worksheets': worksheets
                },
                'properties': properties,
                'statistics': {
                    'sheet_count': sheet_count,
                    'total_rows': total_rows,
                    'total_columns': total_columns
                },
                'processed_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel document: {e}")
            return {"error": f"Failed to process Excel document: {str(e)}"} 